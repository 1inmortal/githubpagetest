#!/usr/bin/env python3
"""
Script para convertir el manual de usuario de Markdown a Word (.docx)
Requiere: pip install python-docx markdown
"""

import os
import sys
from pathlib import Path

def instalar_dependencias():
    """Instala las dependencias necesarias para la conversión"""
    try:
        import docx
        import markdown
        print("✅ Dependencias ya instaladas")
        return True
    except ImportError:
        print("📦 Instalando dependencias necesarias...")
        os.system("pip install python-docx markdown")
        try:
            import docx
            import markdown
            print("✅ Dependencias instaladas correctamente")
            return True
        except ImportError:
            print("❌ Error al instalar dependencias")
            return False

def convertir_md_a_docx(archivo_md, archivo_docx):
    """Convierte un archivo Markdown a Word (.docx)"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import markdown
        import re
        
        # Leer archivo Markdown
        with open(archivo_md, 'r', encoding='utf-8') as f:
            contenido_md = f.read()
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilos
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Crear estilo para títulos
        titulo_style = doc.styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
        titulo_style.font.name = 'Calibri'
        titulo_style.font.size = Pt(18)
        titulo_style.font.bold = True
        titulo_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Crear estilo para subtítulos
        subtitulo_style = doc.styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
        subtitulo_style.font.name = 'Calibri'
        subtitulo_style.font.size = Pt(14)
        subtitulo_style.font.bold = True
        
        # Crear estilo para código
        codigo_style = doc.styles.add_style('Codigo', WD_STYLE_TYPE.PARAGRAPH)
        codigo_style.font.name = 'Consolas'
        codigo_style.font.size = Pt(10)
        
        # Dividir contenido en líneas
        lineas = contenido_md.split('\n')
        
        i = 0
        while i < len(lineas):
            linea = lineas[i].strip()
            
            if not linea:
                i += 1
                continue
            
            # Título principal (#)
            if linea.startswith('# '):
                titulo = linea[2:].strip()
                p = doc.add_paragraph(titulo, style='TituloPrincipal')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            # Subtítulos (##)
            elif linea.startswith('## '):
                subtitulo = linea[3:].strip()
                doc.add_paragraph(subtitulo, style='Subtitulo')
                
            # Sub-subtítulos (###)
            elif linea.startswith('### '):
                sub_subtitulo = linea[4:].strip()
                p = doc.add_paragraph(sub_subtitulo)
                p.style.font.bold = True
                p.style.font.size = Pt(12)
                
            # Lista numerada
            elif re.match(r'^\d+\.', linea):
                item = re.sub(r'^\d+\.\s*', '', linea)
                doc.add_paragraph(item, style='List Number')
                
            # Lista con viñetas
            elif linea.startswith('- ') or linea.startswith('* '):
                item = linea[2:].strip()
                doc.add_paragraph(item, style='List Bullet')
                
            # Código (```)
            elif linea.startswith('```'):
                # Recopilar líneas de código hasta encontrar ```
                codigo_lines = []
                i += 1
                while i < len(lineas) and not lineas[i].strip().startswith('```'):
                    codigo_lines.append(lineas[i])
                    i += 1
                
                if codigo_lines:
                    codigo = '\n'.join(codigo_lines)
                    p = doc.add_paragraph(codigo, style='Codigo')
                    p.style.font.name = 'Consolas'
                    p.style.font.size = Pt(9)
                    
            # Tabla (|)
            elif '|' in linea and not linea.startswith('|---'):
                # Crear tabla
                columnas = [col.strip() for col in linea.split('|') if col.strip()]
                if len(columnas) > 1:
                    tabla = doc.add_table(rows=1, cols=len(columnas))
                    tabla.style = 'Table Grid'
                    
                    # Agregar encabezados
                    hdr_cells = tabla.rows[0].cells
                    for j, col in enumerate(columnas):
                        hdr_cells[j].text = col
                        hdr_cells[j].paragraphs[0].style.font.bold = True
                    
                    # Agregar filas de datos
                    i += 1
                    while i < len(lineas) and '|' in lineas[i] and not lineas[i].strip().startswith('|---'):
                        fila_datos = [col.strip() for col in lineas[i].split('|') if col.strip()]
                        if len(fila_datos) == len(columnas):
                            row = tabla.add_row()
                            for j, dato in enumerate(fila_datos):
                                row.cells[j].text = dato
                        i += 1
                    i -= 1  # Ajustar índice
                    
            # Texto normal
            else:
                # Procesar texto con formato
                texto = linea
                
                # Negrita (**texto**)
                texto = re.sub(r'\*\*(.*?)\*\*', r'\1', texto)
                
                # Código inline (`código`)
                texto = re.sub(r'`([^`]+)`', r'\1', texto)
                
                # Enlaces [texto](url)
                texto = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', texto)
                
                p = doc.add_paragraph(texto)
                
                # Aplicar formato especial si es necesario
                if '**' in linea:
                    for run in p.runs:
                        if '**' in run.text:
                            run.bold = True
                            run.text = run.text.replace('**', '')
                
            i += 1
        
        # Guardar documento
        doc.save(archivo_docx)
        print(f"✅ Manual convertido exitosamente: {archivo_docx}")
        return True
        
    except Exception as e:
        print(f"❌ Error al convertir: {e}")
        return False

def main():
    """Función principal"""
    print("🔄 Convertidor de Manual Markdown a Word")
    print("=" * 50)
    
    # Verificar dependencias
    if not instalar_dependencias():
        return
    
    # Rutas de archivos
    script_dir = Path(__file__).parent
    archivo_md = script_dir / "Manual_Usuario_Sistema_Dron_CLI_Completo.md"
    archivo_docx = script_dir / "Manual_Usuario_Sistema_Dron_CLI_Completo.docx"
    
    # Verificar que existe el archivo Markdown
    if not archivo_md.exists():
        print(f"❌ No se encontró el archivo: {archivo_md}")
        return
    
    print(f"📖 Archivo fuente: {archivo_md}")
    print(f"📄 Archivo destino: {archivo_docx}")
    
    # Convertir
    if convertir_md_a_docx(archivo_md, archivo_docx):
        print("\n🎉 ¡Conversión completada exitosamente!")
        print(f"📁 Ubicación del archivo Word: {archivo_docx}")
        print("\n💡 Puedes abrir el archivo .docx con Microsoft Word o LibreOffice")
    else:
        print("\n❌ Error en la conversión")

if __name__ == "__main__":
    main()
