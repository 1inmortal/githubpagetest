#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Generador de reportes de ciberseguridad (CLI)

Lee CSVs de amenazas y phishing, calcula KPIs (MTTR/MTTD), agrega por año y tipo,
genera artefactos JSON/CSV y (opcional) gráficas en la carpeta reports/.

Uso rápido:
  python scripts/generate_report.py --fast

Requisitos de columnas esperadas:
  Amenazas (threats):
    - Year
    - Response Time (Hours) -> renombrada a Response_Time_Hours
    - Incident Resolution Time (in Hours) -> renombrada a Resolution_Time_Hours
    - Threat Type -> renombrada a Threat_Type
    - Affected Industry -> renombrada a Affected_Industry

  Phishing (opcional):
    - subject -> Subject
    - body -> Body
    - label -> Label (spam/phishing/1 como positivos)
"""

from __future__ import annotations

import argparse
import json
import logging
import os
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd


# -------------------------- Utilidades de logging ---------------------------
def setup_logger(level: int = logging.INFO) -> None:
    logs_dir = Path("logs")
    logs_dir.mkdir(parents=True, exist_ok=True)
    log_file = logs_dir / f"generate_report_{datetime.now().strftime('%Y%m%d')}.log"
    logging.basicConfig(
        level=level,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler(log_file, encoding="utf-8"),
            logging.StreamHandler(),
        ],
    )


# -------------------------- Búsqueda de archivos ---------------------------
def find_csv(filename: str) -> Optional[Path]:
    cwd = Path.cwd()
    candidates = [
        cwd / filename,                # ubicación actual
        cwd / "notebooks" / filename,  # notebooks/
        cwd / "data" / filename,       # data/
        cwd.parent / filename,         # padre (por si se ejecuta desde notebooks/)
        cwd.parent / "data" / filename,
    ]
    for p in candidates:
        if p.exists():
            return p
    # búsqueda recursiva acotada (hasta 4 niveles desde raíz de proyecto)
    root = cwd
    for p in root.rglob(filename):
        try:
            rel_parts = p.relative_to(root).parts
            if len(rel_parts) <= 4:
                return p
        except Exception:
            return p
    return None


# ------------------------------ Carga datos --------------------------------
ThreatsColumnMap = {
    "Response Time (Hours)": "Response_Time_Hours",
    "Incident Resolution Time (in Hours)": "Resolution_Time_Hours",
    "Threat Type": "Threat_Type",
    "Affected Industry": "Affected_Industry",
}

PhishingColumnMap = {
    "subject": "Subject",
    "body": "Body",
    "label": "Label",
}


def load_threats_csv(path: Path) -> pd.DataFrame:
    usecols = None  # permitir columnas extra pero optimizar si es necesario
    df = pd.read_csv(path)
    # renombrar si están las originales
    existing = {k: v for k, v in ThreatsColumnMap.items() if k in df.columns}
    if existing:
        df = df.rename(columns=existing)

    # validar columnas mínimas
    required_any = ["Year", "Response_Time_Hours", "Resolution_Time_Hours"]
    missing = [c for c in required_any if c not in df.columns]
    if missing:
        logging.warning("Faltan columnas en amenazas: %s", ", ".join(missing))

    # coerción numérica segura
    for c in ["Year", "Response_Time_Hours", "Resolution_Time_Hours"]:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    # limpieza de nulos en métricas
    metric_cols = [c for c in ["Response_Time_Hours", "Resolution_Time_Hours"] if c in df.columns]
    if metric_cols:
        before = len(df)
        df = df.dropna(subset=metric_cols)
        logging.info("Amenazas: filas eliminadas por nulos en métricas: %d", before - len(df))

    return df


def load_phishing_csv(path: Path) -> pd.DataFrame:
    df = pd.read_csv(path)
    existing = {k: v for k, v in PhishingColumnMap.items() if k in df.columns}
    if existing:
        df = df.rename(columns=existing)

    for c in ["Subject", "Body"]:
        if c in df.columns:
            df[c] = df[c].astype(str)

    if "Label" in df.columns:
        df["Label"] = df["Label"].astype(str).str.lower()

    # eliminar filas sin texto
    subset = [c for c in ["Subject", "Body"] if c in df.columns]
    if subset:
        before = len(df)
        df = df.dropna(subset=subset)
        logging.info("Phishing: filas sin texto eliminadas: %d", before - len(df))

    return df


# ------------------------------ KPIs y agregados ---------------------------
@dataclass
class ThreatsKPIs:
    overall_mttr: Optional[float]
    overall_mttd: Optional[float]


def compute_threats_kpis(df: pd.DataFrame) -> ThreatsKPIs:
    mttr = df["Resolution_Time_Hours"].mean() if "Resolution_Time_Hours" in df.columns else None
    mttd = df["Response_Time_Hours"].mean() if "Response_Time_Hours" in df.columns else None
    return ThreatsKPIs(overall_mttr=mttr, overall_mttd=mttd)


def segment_by_year_and_type(df: pd.DataFrame) -> Tuple[Optional[pd.Series], Optional[pd.Series], Optional[pd.Series], Optional[pd.Series]]:
    mttr_by_year = df.groupby("Year")["Resolution_Time_Hours"].mean() if set(["Year", "Resolution_Time_Hours"]) <= set(df.columns) else None
    mttd_by_year = df.groupby("Year")["Response_Time_Hours"].mean() if set(["Year", "Response_Time_Hours"]) <= set(df.columns) else None
    mttr_by_type = df.groupby("Threat_Type")["Resolution_Time_Hours"].mean() if set(["Threat_Type", "Resolution_Time_Hours"]) <= set(df.columns) else None
    mttd_by_type = df.groupby("Threat_Type")["Response_Time_Hours"].mean() if set(["Threat_Type", "Response_Time_Hours"]) <= set(df.columns) else None
    return mttr_by_year, mttd_by_year, mttr_by_type, mttd_by_type


def compute_top_categories(df: pd.DataFrame, col: str, n: int = 10) -> Optional[pd.Series]:
    if col in df.columns:
        return df[col].value_counts().head(n)
    return None


def model_phishing_detection(df: pd.DataFrame) -> Dict[str, Optional[float]]:
    if df.empty:
        return {"actual_phishing": 0, "true_positives": 0, "detection_rate": None}

    keywords = [
        "verify", "account", "password", "urgent", "suspend", "confirm", "login",
        "secure", "bank", "credit card", "invoice", "payment", "winner", "update",
    ]

    def detect(text: str) -> int:
        t = (text or "").lower()
        return int(any(k in t for k in keywords))

    if {"Subject", "Body"} <= set(df.columns):
        text_series = df["Subject"].fillna("") + " " + df["Body"].fillna("")
    elif "Body" in df.columns:
        text_series = df["Body"].fillna("")
    elif "Subject" in df.columns:
        text_series = df["Subject"].fillna("")
    else:
        return {"actual_phishing": 0, "true_positives": 0, "detection_rate": None}

    predicted = text_series.apply(detect)
    if "Label" in df.columns:
        actual = df["Label"].astype(str).str.lower().isin(["1", "spam", "phishing"]).astype(int)
    else:
        actual = pd.Series([0] * len(df), index=df.index)

    true_pos = int(((predicted == 1) & (actual == 1)).sum())
    actual_count = int((actual == 1).sum())
    rate = (true_pos / actual_count * 100.0) if actual_count > 0 else None
    return {"actual_phishing": actual_count, "true_positives": true_pos, "detection_rate": rate}


# ------------------------------ Salida/artefactos --------------------------
def ensure_outdir(outdir: Optional[str]) -> Path:
    if outdir:
        out = Path(outdir)
    else:
        out = Path("reports") / datetime.now().strftime("%Y%m%d_%H%M%S")
    out.mkdir(parents=True, exist_ok=True)
    return out


def write_csv(series_or_df: pd.Series | pd.DataFrame, path: Path) -> None:
    if isinstance(series_or_df, pd.Series):
        series_or_df.to_frame(name="value").to_csv(path, index=True)
    else:
        series_or_df.to_csv(path, index=True)


def write_json(obj: dict, path: Path) -> None:
    with path.open("w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)


def try_plot_threats(outdir: Path, df: pd.DataFrame, kpis: ThreatsKPIs, mttr_by_year: Optional[pd.Series], mttd_by_year: Optional[pd.Series], fast: bool) -> None:
    if fast:
        return
    try:
        import seaborn as sns  # type: ignore
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt  # type: ignore

        # Comparación línea base vs proyección
        if kpis.overall_mttd is not None and kpis.overall_mttr is not None:
            metrics = ["MTTD", "MTTR"]
            baseline = [kpis.overall_mttd, kpis.overall_mttr]
            projected = [kpis.overall_mttd * (1 - 0.56), kpis.overall_mttr * (1 - 0.40)]
            plot_data = pd.DataFrame({
                "Métrica": metrics * 2,
                "Tiempo (Horas)": baseline + projected,
                "Estado": ["Línea Base"] * 2 + ["Proyectado"] * 2,
            })
            plt.figure(figsize=(10, 6))
            sns.barplot(x="Métrica", y="Tiempo (Horas)", hue="Estado", data=plot_data, palette="viridis")
            plt.title("Comparación de Métricas: Línea Base vs Proyección")
            plt.tight_layout()
            plt.savefig(outdir / "comparacion_metricas.png", dpi=150)
            plt.close()

        # Evolución anual
        if (mttr_by_year is not None and not mttr_by_year.empty) or (mttd_by_year is not None and not mttd_by_year.empty):
            plt.figure(figsize=(12, 6))
            if mttd_by_year is not None and not mttd_by_year.empty:
                sns.lineplot(x=mttd_by_year.index, y=mttd_by_year.values, marker="o", label="MTTD")
            if mttr_by_year is not None and not mttr_by_year.empty:
                sns.lineplot(x=mttr_by_year.index, y=mttr_by_year.values, marker="o", label="MTTR")
            plt.title("Evolución Anual de MTTR y MTTD")
            plt.tight_layout()
            plt.savefig(outdir / "evolucion_anual.png", dpi=150)
            plt.close()

        # Boxplot por tipo
        if set(["Threat_Type", "Resolution_Time_Hours"]) <= set(df.columns):
            plt.figure(figsize=(14, 8))
            top_types = df["Threat_Type"].value_counts().head(5).index
            subset = df[df["Threat_Type"].isin(top_types)]
            sns.boxplot(x="Threat_Type", y="Resolution_Time_Hours", data=subset, palette="plasma")
            plt.xticks(rotation=45, ha="right")
            plt.yscale("log")
            plt.title("Distribución MTTR por Tipo de Amenaza")
            plt.tight_layout()
            plt.savefig(outdir / "boxplot_mttr_tipo.png", dpi=150)
            plt.close()
    except Exception as e:
        logging.warning("No se pudieron generar gráficas de amenazas: %s", e)


def try_plot_wordcloud(outdir: Path, phishing_df: pd.DataFrame, fast: bool) -> None:
    if fast:
        return
    try:
        from wordcloud import WordCloud  # type: ignore
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt  # type: ignore

        if {"Subject", "Label"} <= set(phishing_df.columns):
            subjects = " ".join(phishing_df.loc[phishing_df["Label"].str.lower().isin(["1", "spam", "phishing"]), "Subject"].astype(str))
            if subjects.strip():
                wc = WordCloud(width=1000, height=500, background_color="black", colormap="Reds").generate(subjects)
                plt.figure(figsize=(14, 7))
                plt.imshow(wc, interpolation="bilinear")
                plt.axis("off")
                plt.title("WordCloud Asuntos Phishing")
                plt.tight_layout()
                plt.savefig(outdir / "wordcloud_phishing.png", dpi=150)
                plt.close()
    except Exception as e:
        logging.warning("No se pudo generar WordCloud: %s", e)


# ---------------------------------- CLI ------------------------------------
def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Generador de reportes de ciberseguridad")
    p.add_argument("--threats", type=str, default=None, help="Ruta del CSV de amenazas")
    p.add_argument("--phishing", type=str, default=None, help="Ruta del CSV de phishing")
    p.add_argument("--from-year", type=int, default=None, help="Filtrar desde año (inclusive)")
    p.add_argument("--to-year", type=int, default=None, help="Filtrar hasta año (inclusive)")
    p.add_argument("--outdir", type=str, default=None, help="Directorio de salida de reportes")
    p.add_argument("--fast", action="store_true", help="Modo rápido: omite gráficas pesadas")
    p.add_argument("--no-plots", action="store_true", help="Desactivar todas las gráficas")
    p.add_argument("--docx", action="store_true", help="Generar informe en Word (.docx)")
    return p.parse_args()


def main() -> None:
    setup_logger()
    args = parse_args()

    outdir = ensure_outdir(args.outdir)
    logging.info("Directorio de salida: %s", outdir)

    # localizar archivos
    threats_path = Path(args.threats) if args.threats else find_csv("Global_Cybersecurity_Threats_2015-2024.csv")
    phishing_path = Path(args.phishing) if args.phishing else find_csv("CEAS_08.csv")

    if threats_path is None or not threats_path.exists():
        logging.error("No se encontró el CSV de amenazas. Proporcione --threats o coloque el archivo en data/ o notebooks/.")
    else:
        logging.info("Amenazas: %s", threats_path)

    if phishing_path is None or not phishing_path.exists():
        logging.warning("No se encontró el CSV de phishing. El reporte continuará sin phishing.")
    else:
        logging.info("Phishing: %s", phishing_path)

    # cargar datasets
    threats_df = pd.DataFrame()
    phishing_df = pd.DataFrame()
    if threats_path and threats_path.exists():
        threats_df = load_threats_csv(threats_path)
        # filtros por año si aplica
        if not threats_df.empty and "Year" in threats_df.columns:
            if args.from_year is not None:
                threats_df = threats_df[threats_df["Year"] >= args.from_year]
            if args.to_year is not None:
                threats_df = threats_df[threats_df["Year"] <= args.to_year]

    if phishing_path and phishing_path.exists():
        phishing_df = load_phishing_csv(phishing_path)

    # cálculos
    report: Dict[str, object] = {"generated_at": datetime.now().isoformat(), "schema_version": 1}

    if not threats_df.empty:
        kpis = compute_threats_kpis(threats_df)
        mttr_by_year, mttd_by_year, mttr_by_type, mttd_by_type = segment_by_year_and_type(threats_df)
        top_threats = compute_top_categories(threats_df, "Threat_Type", 10)
        top_industries = compute_top_categories(threats_df, "Affected_Industry", 10)

        report["threats"] = {
            "overall_mttr": kpis.overall_mttr,
            "overall_mttd": kpis.overall_mttd,
            "mttr_by_year": (mttr_by_year.to_dict() if mttr_by_year is not None else None),
            "mttd_by_year": (mttd_by_year.to_dict() if mttd_by_year is not None else None),
            "mttr_by_type": (mttr_by_type.to_dict() if mttr_by_type is not None else None),
            "mttd_by_type": (mttd_by_type.to_dict() if mttd_by_type is not None else None),
            "top_threat_types": (top_threats.to_dict() if top_threats is not None else None),
            "top_industries": (top_industries.to_dict() if top_industries is not None else None),
        }

        # escribir CSVs
        if mttr_by_year is not None:
            write_csv(mttr_by_year, outdir / "mttr_by_year.csv")
        if mttd_by_year is not None:
            write_csv(mttd_by_year, outdir / "mttd_by_year.csv")
        if mttr_by_type is not None:
            write_csv(mttr_by_type, outdir / "mttr_by_type.csv")
        if mttd_by_type is not None:
            write_csv(mttd_by_type, outdir / "mttd_by_type.csv")
        if top_threats is not None:
            write_csv(top_threats, outdir / "top_threat_types.csv")
        if top_industries is not None:
            write_csv(top_industries, outdir / "top_industries.csv")

        # resumen KPIs
        kpi_df = pd.DataFrame(
            {"metric": ["overall_mttr", "overall_mttd"], "value": [kpis.overall_mttr, kpis.overall_mttd]}
        )
        write_csv(kpi_df, outdir / "kpis.csv")

        # gráficas
        if not args.no_plots:
            try_plot_threats(outdir, threats_df, kpis, mttr_by_year, mttd_by_year, fast=args.fast)

    else:
        logging.warning("No se generaron métricas de amenazas: dataset vacío o no encontrado.")

    if not phishing_df.empty:
        phishing_metrics = model_phishing_detection(phishing_df)
        report["phishing"] = phishing_metrics
        # wordcloud
        if not args.no_plots:
            try_plot_wordcloud(outdir, phishing_df, fast=args.fast)
    else:
        report["phishing"] = None

    # JSON principal
    write_json(report, outdir / "report.json")

    # resumen markdown
    summary_md = [
        f"# Reporte de Ciberseguridad\n",
        f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n",
    ]
    if "threats" in report and report["threats"]:
        t = report["threats"]  # type: ignore
        summary_md.append("## KPIs\n")
        summary_md.append(f"- MTTR: {t.get('overall_mttr', 'N/A')}\n")
        summary_md.append(f"- MTTD: {t.get('overall_mttd', 'N/A')}\n\n")
    if report.get("phishing"):
        p = report["phishing"]  # type: ignore
        summary_md.append("## Phishing\n")
        summary_md.append(f"- Correos phishing reales: {p.get('actual_phishing', 'N/A')}\n")
        summary_md.append(f"- Verdaderos positivos: {p.get('true_positives', 'N/A')}\n")
        summary_md.append(f"- Tasa detección (%): {p.get('detection_rate', 'N/A')}\n")

    (outdir / "report_summary.md").write_text("".join(summary_md), encoding="utf-8")
    logging.info("Reporte generado en: %s", outdir)

    # Documento Word opcional
    if args.docx:
        try:
            from docx import Document  # type: ignore
            from docx.shared import Inches  # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

            doc = Document()
            # Portada
            title = doc.add_heading('Reporte de Ciberseguridad', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(datetime.now().strftime('%Y-%m-%d %H:%M:%S')).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(' ').alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Resumen ejecutivo
            doc.add_heading('Resumen Ejecutivo', level=1)
            if report.get('threats'):
                t = report['threats']  # type: ignore
                p = doc.add_paragraph()
                p.add_run('MTTR: ').bold = True
                p.add_run(str(t.get('overall_mttr', 'N/A')))
                p.add_run('   ')
                p.add_run('MTTD: ').bold = True
                p.add_run(str(t.get('overall_mttd', 'N/A')))
            else:
                doc.add_paragraph('No se cuenta con métricas de amenazas en esta ejecución.')

            # Tablas principales
            def add_table_from_series(title_text: str, series: Optional[pd.Series]):
                if series is None or series.empty:
                    return
                doc.add_heading(title_text, level=2)
                table = doc.add_table(rows=1, cols=2)
                hdr = table.rows[0].cells
                hdr[0].text = 'Clave'
                hdr[1].text = 'Valor'
                for idx, val in series.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(idx)
                    row_cells[1].text = f"{val}"
                doc.add_paragraph('')

            if report.get('threats'):
                # reconstruir series desde dict (si existen)
                t = report['threats']  # type: ignore
                add_table_from_series('MTTR por Año', pd.Series(t.get('mttr_by_year', {})))
                add_table_from_series('MTTD por Año', pd.Series(t.get('mttd_by_year', {})))
                add_table_from_series('MTTR por Tipo', pd.Series(t.get('mttr_by_type', {})))
                add_table_from_series('MTTD por Tipo', pd.Series(t.get('mttd_by_type', {})))
                add_table_from_series('Top Tipos de Amenaza', pd.Series(t.get('top_threat_types', {})))
                add_table_from_series('Top Industrias Afectadas', pd.Series(t.get('top_industries', {})))

            if report.get('phishing'):
                doc.add_heading('Phishing', level=1)
                ph = report['phishing']  # type: ignore
                table = doc.add_table(rows=4, cols=2)
                rows = [
                    ('Correos phishing reales', ph.get('actual_phishing', 'N/A')),
                    ('Verdaderos positivos', ph.get('true_positives', 'N/A')),
                    ('Tasa de detección (%)', ph.get('detection_rate', 'N/A')),
                    ('Notas', 'Heurística basada en palabras clave'),
                ]
                for i, (k, v) in enumerate(rows):
                    r = table.rows[i].cells
                    r[0].text = str(k)
                    r[1].text = str(v)
                doc.add_paragraph('')

            # Anexos de gráficas si existen
            figures = [
                ('comparacion_metricas.png', 'Comparación de Métricas'),
                ('evolucion_anual.png', 'Evolución Anual'),
                ('boxplot_mttr_tipo.png', 'Distribución MTTR por Tipo'),
                ('wordcloud_phishing.png', 'Nube de Palabras (Phishing)'),
            ]
            for fname, caption in figures:
                fpath = outdir / fname
                if fpath.exists():
                    doc.add_heading(caption, level=2)
                    try:
                        doc.add_picture(str(fpath), width=Inches(6.5))
                    except Exception:
                        doc.add_paragraph(f"(No fue posible incrustar la imagen {fname})")
                    doc.add_paragraph('')

            doc_path = outdir / 'reporte_ciberseguridad.docx'
            doc.save(str(doc_path))
            logging.info("Documento Word generado en: %s", doc_path)
        except Exception as e:
            logging.error("No se pudo generar el documento Word: %s", e)


if __name__ == "__main__":
    main()


